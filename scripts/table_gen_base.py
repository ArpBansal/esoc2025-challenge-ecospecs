import os
import docx
import json
import re
from typing import List, Any, Optional, Tuple, Dict, Union
import torch
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches

class BaseTableGenerator:
    """Base class for all table generators."""
    
    def __init__(self, model_name: str = None):
        """Initialize the base table generator."""
        self.model_name = model_name
        self.model = None
        self.tokenizer = None
        self.pipeline = None
        self.system_prompt = """You are a helpful assistant that generates tables based on the provided prompt.
You will receive a prompt and you need to generate a table in text format.
Return ONLY a valid JSON array of arrays like: [[value1, value2, ...], [value1, value2, ...], ...]
Each inner array represents one row of data.
Do not include any explanation, markdown formatting, or code blocks."""

    def _setup_model(self):
        """Setup the model and tokenizer. To be implemented by subclasses."""
        raise NotImplementedError("Subclasses must implement _setup_model")
    
    def _format_prompt(self, prompt: str, row_headers: Optional[List[str]] = None,
                      column_headers: Optional[List[str]] = None,
                      num_rows: int = 5, num_cols: int = 3) -> str:
        """Format prompt for table generation."""
        message = f"Generate content for a table about: {prompt}\n\n"
        
        if row_headers and column_headers:
            message += f"Row headers: {', '.join(row_headers)}\n"
            message += f"Column headers: {', '.join(column_headers)}\n"
            message += "Generate exactly one value for each cell in the table.\n"
        else:
            if row_headers:
                message += f"Row headers: {', '.join(row_headers)}\n"
                message += f"Generate {num_cols} columns of appropriate content.\n"
            elif column_headers:
                message += f"Column headers: {', '.join(column_headers)}\n"
                message += f"Generate {num_rows} rows of appropriate content.\n"
            else:
                message += f"Generate a {num_rows}x{num_cols} table with appropriate content.\n"
        
        message += "\nReturn your response only as a JSON array of arrays (rows)"
        message += "\nExample format: [[\"cell1\", \"cell2\"], [\"cell3\", \"cell4\"]]"
        return message

    def _extract_json_array(self, text: str) -> str:
        """Extract JSON array from text with improved parsing."""
        # Remove code block markers
        text = re.sub(r'```(?:json)?\s*|\s*```', '', text)
        
        # Find outermost brackets containing valid JSON array
        bracket_start = text.find('[')
        if bracket_start == -1:
            return '[[]]'  # Return valid empty JSON if no array found
            
        # Count brackets to find matching closing bracket
        open_count = 0
        for i in range(bracket_start, len(text)):
            if text[i] == '[':
                open_count += 1
            elif text[i] == ']':
                open_count -= 1
                if open_count == 0:
                    # Found matching brackets, extract content
                    return text[bracket_start:i+1]
        
        # If no properly matched brackets found
        return '[[]]'

    def _extract_json_object(self, text: str) -> str:
        """Extract JSON object from text."""
        # Remove code block markers
        text = re.sub(r'```(?:json)?\s*|\s*```', '', text)
        
        # Find outermost braces containing valid JSON object
        brace_start = text.find('{')
        if brace_start == -1:
            return '{"row_headers":[],"column_headers":[],"num_rows":5,"num_cols":3}'
            
        # Count braces to find matching closing brace
        open_count = 0
        for i in range(brace_start, len(text)):
            if text[i] == '{':
                open_count += 1
            elif text[i] == '}':
                open_count -= 1
                if open_count == 0:
                    # Found matching braces, extract content
                    return text[brace_start:i+1]
        
        # If no properly matched braces found
        return '{"row_headers":[],"column_headers":[],"num_rows":5,"num_cols":3}'

    def _parse_json_content(self, json_text: str, rows: int, cols: int) -> List[List[str]]:
        """Parse JSON content into a properly dimensioned table."""
        try:
            # Try to parse as JSON
            table_data = json.loads(json_text)
            
            # Ensure proper structure - should be list of lists
            if not isinstance(table_data, list):
                table_data = [[]]
            elif table_data and not isinstance(table_data[0], list):
                # If we got a flat array, convert it to 2D
                table_data = [table_data]
            
            # Ensure we have correct dimensions
            result = []
            for i in range(min(rows, len(table_data))):
                row = table_data[i]
                if not isinstance(row, list):
                    row = [str(row)]  # Convert non-list rows to list
                
                # Ensure each row has correct number of columns
                new_row = []
                for j in range(cols):
                    if j < len(row):
                        # Convert any value to string
                        new_row.append(str(row[j]) if row[j] is not None else "")
                    else:
                        new_row.append("")  # Fill missing columns
                
                result.append(new_row)
            
            # Add any missing rows
            while len(result) < rows:
                result.append([""] * cols)
            
            return result
            
        except json.JSONDecodeError as e:
            print(f"Failed to parse JSON response: {e}")
            print(f"Problematic JSON text: {json_text}")
            return [[""] * cols for _ in range(rows)]

    def save_table_to_file(self, table_data: List[List[str]], output_file: str,
                         intro_text: Optional[str] = None, 
                         table_title: Optional[str] = None,
                         as_pdf: bool = False) -> str:
        """Save a table to a Word document with optional introduction and title.s
        Args:
            table_data (List[List[str]]): Table data to save.
            output_file (str): Path to save the document.
            intro_text (Optional[str]): Introduction text to add.
            table_title (Optional[str]): Title for the table.
            as_pdf (bool): Whether to save as PDF. False = docx generated"""
        doc = docx.Document()

        # Add title if provided
        if table_title:
            heading = doc.add_heading(table_title, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        if intro_text:
            para = doc.add_paragraph(intro_text)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            doc.add_paragraph()  # Add space after intro
        
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0

        if rows > 0 and cols > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # Auto-fit table to page width
            table.autofit = True
            
            # Set cell contents and formatting
            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    if j < cols:
                        cell_obj = table.cell(i, j)
                        cell_obj.text = str(cell) if cell is not None else ""
                        
                        # Header formatting (first row and first column)
                        if i == 0 or j == 0:
                            for paragraph in cell_obj.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

        # Save as DOCX
        saved_path = output_file
        doc.save(output_file)
        
        # Convert to PDF if requested
        if as_pdf:
            try:
                base_name = os.path.splitext(output_file)[0]
                pdf_path = f"{base_name}.pdf"
                
                from scripts.utils import convert_docx_to_pdf
                convert_docx_to_pdf(saved_path, output_file)
                os.remove(saved_path)  # Remove temp DOCX
                print(f"Table saved to {output_file}")
            except ImportError:
                print("Warning: docx2pdf package not installed. Saved as DOCX instead.")
                print(f"Table saved to {output_file}")
        else:
            print(f"Table saved to {output_file}")
            
        return saved_path


class EnhancedHuggingFaceTableGenerator(BaseTableGenerator):
    """Enhanced table generator using Hugging Face transformers with both auto and manual header options."""
    
    def __init__(self, model_name: str = "Qwen/Qwen2.5-7B-Instruct-AWQ"):
        """Initialize HuggingFace table generator."""
        super().__init__(model_name)
        self._setup_model()
        
    def _setup_model(self):
        """Setup the model and tokenizer."""
        from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline
        
        try:
            print(f"Loading model: {self.model_name}")
            self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
            self.model = AutoModelForCausalLM.from_pretrained(
                self.model_name,
                device_map="cuda" if torch.cuda.is_available() else "cpu",
                torch_dtype=torch.float16
            )
            
            self.pipeline = pipeline(
                model=self.model,
                tokenizer=self.tokenizer,
                task="text-generation",
                device_map="cuda" if torch.cuda.is_available() else "cpu",
                torch_dtype=torch.float16,
            )
            print("Model loaded successfully")
        except Exception as e:
            print(f"Error loading model: {e}")
            raise
            
    def _format_messages(self, user_content: str) -> List[Dict[str, str]]:
        """Format messages for specific model types."""
        
        return [
            {"role": "system", "content": self.system_prompt},
            {"role": "user", "content": user_content}
        ]
    
    def generate_table_structure(self, prompt: str) -> Dict:
        """Generate table structure including headers."""
        structure_prompt = f"""Create an appropriate table structure for: {prompt}
        
Generate row headers and column headers that would make sense for this type of table.
Return a JSON object with the keys:
- 'row_headers': an array of row headers
- 'column_headers': an array of column headers
- 'num_rows': suggested number of rows (between 3-10)
- 'num_cols': suggested number of columns (between 3-8)

IMPORTANT: Only return valid JSON with no additional text."""

        messages = self._format_messages(structure_prompt)
        
        try:
            response = self.pipeline(
                messages, 
                max_new_tokens=1024,
                do_sample=True,
                temperature=0.7,
                top_p=0.95
            )
            content = response[0]['generated_text'][-1]['content']
            print(content)
            
            # Extract JSON object
            clean_content = self._extract_json_object(content)
            structure = json.loads(clean_content)
            
            # Validate structure has required fields
            if not all(key in structure for key in ['row_headers', 'column_headers']):
                raise ValueError("Generated structure missing required fields")
                
            return structure
        except Exception as e:
            print(f"Error generating table structure: {e}")
            # Return default structure
            return {
                "row_headers": ["Row 1", "Row 2", "Row 3", "Row 4", "Row 5"],
                "column_headers": ["Column 1", "Column 2", "Column 3"],
                "num_rows": 5,
                "num_cols": 3
            }
    
    def generate_intro_paragraph(self, prompt: str, table_structure: Dict) -> str:
        """Generate introduction paragraph for the table."""
        intro_prompt = f"""Write a brief introduction paragraph for a table about: {prompt}
        
The table has the following structure:
- Row headers: {table_structure.get('row_headers', [])}
- Column headers: {table_structure.get('column_headers', [])}

The introduction should be 2-4 sentences, professional in tone, and explain what information the table contains and its relevance.
Be concise and direct."""

        messages = self._format_messages(intro_prompt)
        
        try:
            response = self.pipeline(
                messages, 
                max_new_tokens=512,
                do_sample=True,
                temperature=0.7
            )
            content =response[0]['generated_text'][-1]['content']
            print(content)
            
            # Extract just the intro paragraph (removing any json, code blocks, etc.)
            # Remove any JSON, markdown code blocks
            clean_content = re.sub(r'```.*?```', '', content, flags=re.DOTALL)
            clean_content = re.sub(r'\{.*?\}', '', clean_content, flags=re.DOTALL)
            
            # Get the last few paragraphs which should be the intro
            paragraphs = [p for p in clean_content.split('\n\n') if p.strip()]
            intro = paragraphs[-1] if paragraphs else ""
            
            return intro.strip()
        except Exception as e:
            print(f"Error generating introduction: {e}")
            return f"This table presents information about {prompt}."
        
    def generate_table_content(self, prompt: str,
                              row_headers: Optional[List[str]] = None,
                              column_headers: Optional[List[str]] = None,
                              num_rows: int = 5,
                              num_cols: int = 3) -> List[List[str]]:
        """Generate content for a table with either provided or auto-generated headers."""
        rows = len(row_headers) if row_headers else num_rows
        cols = len(column_headers) if column_headers else num_cols
        
        content_prompt = self._format_prompt(prompt, row_headers, column_headers, rows, cols)
        messages = self._format_messages(content_prompt)
        
        try:
            response = self.pipeline(
                messages, 
                max_new_tokens=4096,
                do_sample=True,
                temperature=0.7
            )
            content = response[0]['generated_text'][-1]['content']
            print(content)
            # Clean and extract JSON
            clean_content = self._extract_json_array(content)
            table_data = self._parse_json_content(clean_content, rows, cols)
            
            return table_data
        except Exception as e:
            print(f"Error generating table content: {e}")
            # Return empty table as fallback
            return [[""] * cols for _ in range(rows)]
    
    def generate_complete_table_auto(self, prompt: str) -> Tuple[Dict, List[List[str]], str]:
        """Generate complete table with auto-generated structure, content and intro."""
        # First, generate table structure
        print("Generating table structure...")
        structure = self.generate_table_structure(prompt)
        
        row_headers = structure.get('row_headers', [])
        column_headers = structure.get('column_headers', [])
        num_rows = len(row_headers)
        num_cols = len(column_headers)
        
        # Next, generate table content
        print("Generating table content...")
        content = self.generate_table_content(
            prompt, 
            row_headers, 
            column_headers,
            num_rows,
            num_cols
        )
        
        # Generate introduction paragraph
        print("Generating introduction paragraph...")
        intro = self.generate_intro_paragraph(prompt, structure)
        
        return structure, content, intro
    
    def generate_complete_table_manual(self, prompt: str, 
                                     row_headers: Optional[List[str]] = None,
                                     column_headers: Optional[List[str]] = None,
                                     num_rows: int = None,
                                     num_cols: int = None) -> Tuple[Dict, List[List[str]], str]:
        """Generate complete table with manually specified headers."""
        # Create structure from provided headers
        if not row_headers and not column_headers:
            # If neither is provided, use auto generation
            return self.generate_complete_table_auto(prompt)
        
        structure = {
            "row_headers": row_headers or [],
            "column_headers": column_headers or [],
            "num_rows": len(row_headers) if row_headers else (num_rows or 5),
            "num_cols": len(column_headers) if column_headers else (num_cols or 3)
        }
        
        # Generate table content based on provided headers
        print("Generating table content with provided headers...")
        content = self.generate_table_content(
            prompt, 
            structure["row_headers"], 
            structure["column_headers"],
            structure["num_rows"],
            structure["num_cols"]
        )
        
        # Generate introduction paragraph
        print("Generating introduction paragraph...")
        intro = self.generate_intro_paragraph(prompt, structure)
        
        return structure, content, intro
    
    def create_and_save_table(self, prompt: str, output_file: str,
                            row_headers: Optional[List[str]] = None,
                            column_headers: Optional[List[str]] = None,
                            num_rows: int = None,
                            num_cols: int = None,
                            as_pdf: bool = False,
                            table_title: Optional[str] = None) -> str:
        """Generate and save a table with either auto or manual headers."""
        # Determine if we're using manual or auto mode
        if row_headers is not None or column_headers is not None:
            # Manual mode
            structure, content, intro = self.generate_complete_table_manual(
                prompt, row_headers, column_headers, num_rows, num_cols
            )
        else:
            # Auto mode
            structure, content, intro = self.generate_complete_table_auto(prompt)
        
        row_headers = structure.get('row_headers', [])
        column_headers = structure.get('column_headers', [])
        
        # Format table with headers
        formatted_table = [[""] + column_headers]  # First cell empty, then column headers
        
        for i, header in enumerate(row_headers):
            if i < len(content):
                formatted_table.append([header] + content[i])
            else:
                formatted_table.append([header] + [""] * len(column_headers))
        
        # Generate title if not provided
        if not table_title:
            table_title = f"Table: {prompt.capitalize()}"
            
        # Save to file
        saved_path = self.save_table_to_file(
            formatted_table,
            output_file,
            intro_text=intro,
            table_title=table_title,
            as_pdf=as_pdf
        )
        
        return saved_path


# Example usage
if __name__ == "__main__":
    # Auto-generated headers example
    generator = EnhancedHuggingFaceTableGenerator(model_name="Qwen/Qwen2.5-7B-Instruct-AWQ")
    
    # Example 1: Auto-generate both row and column headers
    generator.create_and_save_table(
        prompt="Comparison of popular programming languages",
        output_file="auto_programming_languages.docx"
    )
    
    # Example 2: Manually specify row headers
    row_headers = ["Python", "JavaScript", "Java", "C++", "Rust"]
    generator.create_and_save_table(
        prompt="Programming language features and use cases",
        output_file="manual_rows_programming_languages.docx",
        row_headers=row_headers
    )
    
    # Example 3: Manually specify column headers
    column_headers = ["Learning Curve", "Performance", "Community", "Use Cases"]
    generator.create_and_save_table(
        prompt="Programming language comparisons",
        output_file="manual_cols_programming_languages.docx",
        column_headers=column_headers
    )
    
    # Example 4: Manually specify both row and column headers
    row_headers = ["Web Development", "Data Science", "Mobile Apps", "System Programming"]
    column_headers = ["Best Languages", "Popular Frameworks", "Learning Resources"]
    generator.create_and_save_table(
        prompt="Programming domains and technologies",
        output_file="manual_both_programming_languages.docx",
        row_headers=row_headers,
        column_headers=column_headers
    )