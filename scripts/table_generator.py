import os
import docx
import json
import re
from typing import List, Optional
from transformers import pipeline, AutoModelForCausalLM, AutoTokenizer
import torch

class TableGenerator:
    def __init__(self, model_path: str):
        self.model = AutoModelForCausalLM.from_pretrained(
            model_path,
            device_map="cuda" if torch.cuda.is_available() else "cpu",
            torch_dtype=torch.float16
        )
        self.tokenizer = AutoTokenizer.from_pretrained(model_path)

        self.pipe = pipeline(
            model=self.model,
            tokenizer=self.tokenizer,
            task="text-generation",
            device_map="cuda" if torch.cuda.is_available() else "cpu",
            torch_dtype=torch.float16,
        )
        
        self.SYSTEM_PROMPT = """You are a helpful assistant that generates tables based on the provided prompt.
You will receive a prompt and you need to generate a table in text format.
Return ONLY a valid JSON array of arrays like: [[value1, value2, ...], [value1, value2, ...], ...]
Each inner array represents one row of data.
Do not include any explanation, markdown formatting, or code blocks."""
        
    def generate_table_content(self, prompt: str, max_new_tokens: int = 2048,
                            row_headers: Optional[List[str]] = None,
                            column_headers: Optional[List[str]] = None,
                            num_rows: int = 5,
                            num_cols: int = 3) -> List[List[str]]:
        
        rows = len(row_headers) if row_headers else num_rows
        cols = len(column_headers) if column_headers else num_cols
        
        # Create a simplified, explicit prompt
        message = f"Generate content for a table about: {prompt}\n\n"
        
        if row_headers and column_headers:
            message += f"Row headers: {', '.join(row_headers)}\n"
            message += f"Column headers: {', '.join(column_headers)}\n"
            message += "Generate exactly one value for each cell in the table.\n"
        else:
            if row_headers:
                message += f"Row headers: {', '.join(row_headers)}\n"
                message += f"Generate {cols} columns of appropriate content.\n"
            elif column_headers:
                message += f"Column headers: {', '.join(column_headers)}\n"
                message += f"Generate {rows} rows of appropriate content.\n"
            else:
                message += f"Generate a {rows}x{cols} table with appropriate content.\n"
        
        message += "\nFormat your response ONLY as a JSON array of arrays. No explanation text or markdown formatting."
        message += "\nExample format: [[\"cell1\", \"cell2\"], [\"cell3\", \"cell4\"]]"
        message += "\nRemember to only "
        messages = [
            {"role": "system", "content": self.SYSTEM_PROMPT},
            {"role": "user", "content": message}
        ]
        
        try:
            response = self.pipe(messages, max_new_tokens=max_new_tokens)
            content = response[0]['generated_text'][-1]['content']
            print("mannual content:\n", content)

            clean_content = self._extract_json_array(content)
            print("clean content:\n", clean_content)
            
            table_data = self._parse_json_content(clean_content, rows, cols)
            print("table data:\n", table_data)
            
            return table_data
        except Exception as e:
            print(f"Error generating table content: {e}")
            return [[""] * cols for _ in range(rows)]

    def _extract_json_array(self, text: str) -> str:
        """Extract JSON array from text with improved parsing"""
        text = re.sub(r'```(?:json)?\s*|\s*```', '', text)
        
        bracket_start = text.find('[')
        if bracket_start == -1:
            return '[[]]'
            
        open_count = 0
        for i in range(bracket_start, len(text)):
            if text[i] == '[':
                open_count += 1
            elif text[i] == ']':
                open_count -= 1
                if open_count == 0:
                    # Found matching brackets, extract content
                    return text[bracket_start:i+1]
        
        return '[[]]'

    def _parse_json_content(self, json_text: str, rows: int, cols: int) -> List[List[str]]:
        """Parse JSON content into a properly dimensioned table."""
        try:
            table_data = json.loads(json_text)
            
            if not isinstance(table_data, list):
                table_data = [[]]
            elif table_data and not isinstance(table_data[0], list):
                table_data = [table_data]
            
            result = []
            actual_cols = max([len(row) if isinstance(row, list) else 1 for row in table_data]) if table_data else cols
            
            for i in range(min(rows, len(table_data))):
                row = table_data[i]
                if not isinstance(row, list):
                    row = [str(row)]
                
                # For each row, use the actual number of columns in that row
                new_row = []
                for j in range(min(actual_cols, len(row))):
                    new_row.append(str(row[j]) if row[j] is not None else "")
                
                result.append(new_row)
            
            while len(result) < rows:
                result.append([""] * actual_cols if result else [""] * cols)
            
            return result
            
        except json.JSONDecodeError as e:
            print(f"Failed to parse JSON response: {e}")
            print(f"Problematic JSON text: {json_text}")
            return [[""] * cols for _ in range(rows)]        
    def save_table_to_docx(self, table_data: List[List[str]], output_file: str,
                           intro_text: Optional[str] = None) -> None:
        """Save a table to a Word document with optional introduction paragraph."""
        doc = docx.Document()

        if intro_text:
            doc.add_paragraph(intro_text)
            doc.add_paragraph()
        
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0

        if rows > 0 and cols > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'

            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    if j < cols:
                        table.cell(i, j).text = str(cell) if cell is not None else ""

        doc.save(output_file)
        print(f"Table saved to {output_file}")

    def save_table_to_file(self, table_data: List[List[str]], output_file: str,
                     intro_text: Optional[str] = None, as_pdf: bool = False) -> None:
        """Save a table to a Word document or PDF file with optional introduction paragraph."""
        doc = docx.Document()

        if intro_text:
            doc.add_paragraph(intro_text)
            doc.add_paragraph()
        
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0

        if rows > 0 and cols > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'

            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    if j < cols:
                        table.cell(i, j).text = str(cell) if cell is not None else ""

        temp_docx = output_file
        if as_pdf:
            base_name = os.path.splitext(output_file)[0]
            temp_docx = f"{base_name}_temp.docx"

        doc.save(temp_docx)
        
        if as_pdf:
            try:
                from scripts.utils import convert_docx_to_pdf
                convert_docx_to_pdf(temp_docx, output_file)
                os.remove(temp_docx)
                print(f"Table saved to {output_file}")
            except ImportError:
                print("Warning: docx2pdf package not installed. Saving as DOCX instead.")
                print(f"Table saved to {temp_docx}")
        else:
            print(f"Table saved to {output_file}")