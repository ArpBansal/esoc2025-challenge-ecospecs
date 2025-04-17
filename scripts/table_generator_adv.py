import os
import docx
import json
import re
from typing import List, Any, Optional, Tuple, Dict
from transformers import pipeline, AutoModelForCausalLM, AutoTokenizer
import torch
from scripts.table_generator import TableGenerator

class AdvancedTableGenerator:
    def __init__(self, model_path: str):
        self.model = AutoModelForCausalLM.from_pretrained(
            model_path,
            device_map="cuda" if torch.cuda.is_available() else "cpu",
            torch_dtype=torch.float16
        )
        self.tokenizer = AutoTokenizer.from_pretrained(model_path)

        self.pipe = pipeline(
            model=self.model,
            tokenizer=self.tokenizer,
            task="text-generation",
            device_map="cuda" if torch.cuda.is_available() else "cpu",
            torch_dtype=torch.float16,
        )
        
        self.SYSTEM_PROMPT = """You are a table generation assistant that creates table content in JSON format.
        Return ONLY valid JSON with no additional text or formatting.
        For table structure requests, return a JSON object with 'row_headers', 'column_headers', and 'data' keys.
        For table content requests, return a JSON array of arrays representing rows of data."""
        
    def generate_table_structure(self, prompt: str, max_new_tokens: int = 2048) -> Dict:
        """Generate both row and column headers for a table based on a prompt"""
        message = f"""Create an appropriate table structure for: {prompt}
        
        Generate row headers and column headers that would make sense for this type of table.
        Return a JSON object with the keys:
        - 'row_headers': an array of row headers
        - 'column_headers': an array of column headers
        - 'num_rows': suggested number of rows (between 3-10)
        - 'num_cols': suggested number of columns (between 3-8)

        Example format:
        {{
        "row_headers": ["Item 1", "Item 2", "Item 3"],
        "column_headers": ["Property 1", "Property 2", "Property 3"],
        "num_rows": 3,
        "num_cols": 3
        }}

        Only return valid JSON with no additional text."""
        
        messages = [
            {"role": "system", "content": self.SYSTEM_PROMPT},
            {"role": "user", "content": message}
        ]
        
        try:
            response = self.pipe(messages, max_new_tokens=max_new_tokens)
            content = response[0]['generated_text'][-1]['content']
            
            clean_content = self._extract_json_object(content)
            structure = json.loads(clean_content)
            
            if not all(key in structure for key in ['row_headers', 'column_headers']):
                raise ValueError("Generated structure missing required fields")
                
            return structure
        except Exception as e:
            print(f"Error generating table structure: {e}")

            return {
                "row_headers": ["Row 1", "Row 2", "Row 3", "Row 4", "Row 5"],
                "column_headers": ["Column 1", "Column 2", "Column 3"],
                "num_rows": 5,
                "num_cols": 3
            }
    
    def generate_intro_paragraph(self, prompt: str, table_structure: Dict) -> str:
        """Generate an introduction paragraph for the table"""
        message = f"""Write a brief introduction paragraph for a table about: {prompt}
        
The table has the following structure:
- Row headers: {table_structure.get('row_headers', [])}
- Column headers: {table_structure.get('column_headers', [])}

The introduction should be 2-4 sentences, professional in tone, and explain what information the table contains and its relevance."""
        
        messages = [
            {"role": "system", "content": "You are a professional technical writer who creates clear, concise document introductions."},
            {"role": "user", "content": message}
        ]
        
        try:
            response = self.pipe(messages, max_new_tokens=512)
            content = response[0]['generated_text'][-1]['content']
            return content.strip()
        except Exception as e:
            print(f"Error generating introduction: {e}")
            return f"This table presents information about {prompt}."
        
    def generate_table_content(self, prompt: str, max_new_tokens: int = 2048,
                            row_headers: Optional[List[str]] = None,
                            column_headers: Optional[List[str]] = None,
                            num_rows: int = 5,
                            num_cols: int = 3) -> List[List[str]]:
        """Generate content for a table based on row/column headers"""

        rows = len(row_headers) if row_headers else num_rows
        cols = len(column_headers) if column_headers else num_cols
        
        message = f"Generate content for a table about: {prompt}\n\n"
        
        if row_headers and column_headers:
            message += f"Row headers: {', '.join(row_headers)}\n"
            message += f"Column headers: {', '.join(column_headers)}\n"
            message += "Generate exactly one value for each cell in the table.\n"
        else:
            if row_headers:
                message += f"Row headers: {', '.join(row_headers)}\n"
                message += f"Generate {cols} columns of appropriate content.\n"
            elif column_headers:
                message += f"Column headers: {', '.join(column_headers)}\n"
                message += f"Generate {rows} rows of appropriate content.\n"
            else:
                message += f"Generate a {rows}x{cols} table with appropriate content.\n"
        
        message += "\nReturn ONLY a JSON array of arrays representing rows of data."
        message += "\nExample format: [[\"cell1\", \"cell2\"], [\"cell3\", \"cell4\"]]"
        
        messages = [
            {"role": "system", "content": self.SYSTEM_PROMPT},
            {"role": "user", "content": message}
        ]
        
        try:
            response = self.pipe(messages, max_new_tokens=max_new_tokens)
            content = response[0]['generated_text'][-1]['content']
            
            clean_content = self._extract_json_array(content)
            table_data = self._parse_json_content(clean_content, rows, cols)
            
            return table_data
        except Exception as e:
            print(f"Error generating table content: {e}")
            return [[""] * cols for _ in range(rows)]

    def generate_complete_table(self, prompt: str, max_new_tokens: int = 4096) -> Tuple[Dict, List[List[str]], str]:
        """Generate a complete table including structure, content, and introduction paragraph"""

        print("Generating table structure...")
        structure = self.generate_table_structure(prompt, max_new_tokens)
        
        row_headers = structure.get('row_headers', [])
        column_headers = structure.get('column_headers', [])
        num_rows = len(row_headers)
        num_cols = len(column_headers)
        
        print("Generating table content...")
        content = self.generate_table_content(
            prompt, 
            max_new_tokens, 
            row_headers, 
            column_headers,
            num_rows,
            num_cols
        )
        
        print("Generating introduction paragraph...")
        intro = self.generate_intro_paragraph(prompt, structure)
        
        return structure, content, intro

    def _extract_json_array(self, text: str) -> str:
        """Extract JSON array from text"""
        text = re.sub(r'```(?:json)?\s*|\s*```', '', text)
        
        bracket_start = text.find('[')
        if bracket_start == -1:
            return '[[]]'
            
        open_count = 0
        for i in range(bracket_start, len(text)):
            if text[i] == '[':
                open_count += 1
            elif text[i] == ']':
                open_count -= 1
                if open_count == 0:
                    return text[bracket_start:i+1]
        
        # If no properly matched brackets found
        return '[[]]'
        
    def _extract_json_object(self, text: str) -> str:
        """Extract JSON object from text"""
        text = re.sub(r'```(?:json)?\s*|\s*```', '', text)
        
        brace_start = text.find('{')
        if brace_start == -1:
            return '{"row_headers":[],"column_headers":[],"num_rows":5,"num_cols":3}'
            
        open_count = 0
        for i in range(brace_start, len(text)):
            if text[i] == '{':
                open_count += 1
            elif text[i] == '}':
                open_count -= 1
                if open_count == 0:
                    return text[brace_start:i+1]
        
        # If no properly matched braces found
        return '{"row_headers":[],"column_headers":[],"num_rows":5,"num_cols":3}'

    def _parse_json_content(self, json_text: str, rows: int, cols: int) -> List[List[str]]:
        """Parse JSON content into a properly dimensioned table"""
        try:
            table_data = json.loads(json_text)
            
            if not isinstance(table_data, list):
                table_data = [[]]
            elif table_data and not isinstance(table_data[0], list):
                table_data = [table_data]
            
            result = []
            for i in range(min(rows, len(table_data))):
                row = table_data[i]
                if not isinstance(row, list):
                    row = [str(row)]
                
                new_row = []
                for j in range(cols):
                    if j < len(row):
                        new_row.append(str(row[j]) if row[j] is not None else "")
                    else:
                        new_row.append("")
                
                result.append(new_row)
            
            while len(result) < rows:
                result.append([""] * cols)
            
            return result
            
        except json.JSONDecodeError as e:
            print(f"Failed to parse JSON response: {e}")
            print(f"Problematic JSON text: {json_text}")
            return [[""] * cols for _ in range(rows)]
            
    def save_table_to_docx(self, table_data: List[List[str]], output_file: str,
                           intro_text: Optional[str] = None,
                           table_title: Optional[str] = None) -> None:
        """Save a table to a Word document with optional introduction paragraph and title"""
        doc = docx.Document()

        if table_title:
            doc.add_heading(table_title, level=1)
            
        if intro_text:
            doc.add_paragraph(intro_text)
            doc.add_paragraph()
        
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0

        if rows > 0 and cols > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'

            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    if j < cols:
                        table.cell(i, j).text = str(cell) if cell is not None else ""

        doc.save(output_file)
        print(f"Table saved to {output_file}")

    def save_table_to_file(self, table_data: List[List[str]], output_file: str,
                     intro_text: Optional[str] = None, 
                     table_title: Optional[str] = None,
                     as_pdf: bool = False) -> None:
        """Save a table to a Word document or PDF file with optional introduction paragraph"""
        doc = docx.Document()

        if table_title:
            doc.add_heading(table_title, level=1)
            
        if intro_text:
            doc.add_paragraph(intro_text)
            doc.add_paragraph()
        
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0

        if rows > 0 and cols > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'

            for i, row in enumerate(table_data):
                for j, cell in enumerate(row):
                    if j < cols:
                        table.cell(i, j).text = str(cell) if cell is not None else ""

        temp_docx = output_file
        if as_pdf:
            base_name = os.path.splitext(output_file)[0]
            temp_docx = f"{base_name}_temp.docx"

        doc.save(temp_docx)
        
        if as_pdf:
            try:
                from scripts.utils import convert_docx_to_pdf
                convert_docx_to_pdf(temp_docx, output_file)
                os.remove(temp_docx)
                print(f"Table saved to {output_file}")
            except ImportError:
                print("Warning: docx2pdf package not installed. Saving as DOCX instead.")
                print(f"Table saved to {temp_docx}")
        else:
            print(f"Table saved to {output_file}")

    def create_and_save_complete_table(self, prompt: str, output_file: str, 
                                    as_pdf: bool = False,
                                    table_title: Optional[str] = None) -> None:
        """Generate a complete table and save it to a file in one step"""
        structure, content, intro = self.generate_complete_table(prompt)
        print("structure:", structure)
        print("content:", content)
        print("intro:", intro)
        
        row_headers = structure.get('row_headers', [])
        column_headers = structure.get('column_headers', [])
        
        formatted_table = [[""] + column_headers]
        
        for i, header in enumerate(row_headers):
            if i < len(content):
                formatted_table.append([header] + content[i])
            else:
                formatted_table.append([header] + [""] * len(column_headers))
        
        if not table_title:
            table_title = f"Table: {prompt.capitalize()}"
            
        self.save_table_to_file(
            formatted_table,
            output_file,
            intro_text=intro,
            table_title=table_title,
            as_pdf=as_pdf
        )
        
        return formatted_table, intro