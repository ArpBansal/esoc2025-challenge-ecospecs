import docx
import numpy as np
from typing import List, Any


def parse_tables_from_docx(file_path: str) -> List[List[List[Any]]]:
    """
    Parse all tables from a Word document.
    
    Args:
        file_path (str): Path to the Word document
        
    Returns:
        List[List[List[Any]]]: List of tables, where each table is a 2D array
    """
    document = docx.Document(file_path)
    tables = []
    
    for table in document.tables:
        parsed_table = []
        for row in table.rows:
            parsed_row = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                parsed_row.append(cell_text)
            parsed_table.append(parsed_row)
        tables.append(parsed_table)
    
    return tables

def get_table_dimensions(tables: List[List[List[Any]]]) -> List[tuple]:
    """
    Get dimensions of each table in the list.
    
    Args:
        tables (List[List[List[Any]]]): List of tables
        
    Returns:
        List[tuple]: List of (rows, columns) for each table
    """
    
    dimensions = []
    for table in tables:
        num_rows = len(table)
        num_cols = len(table[0]) if num_rows > 0 else 0
        dimensions.append((num_rows, num_cols))
    return dimensions

# Although not needed in this script, we can import the following for future use:
def parse_table(file_path: str, table_index: int = 0) -> np.ndarray:
    """
    Parse a specific table from a Word document.
    
    Args:
        file_path (str): Path to the Word document
        table_index (int): Index of the table to parse (default is 0)
        
    Returns:
        np.ndarray: Parsed table as a NumPy array
    """
    document = docx.Document(file_path)
    tables = document.tables
    
    if table_index >= len(tables):
        raise IndexError(f"Table index {table_index} out of range. Document contains {len(tables)} tables.")
    
    table = tables[table_index]
    parsed_table = []
    
    for row in table.rows:
        parsed_row = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            parsed_row.append(cell_text)
        parsed_table.append(parsed_row)
    
    return np.array(parsed_table)