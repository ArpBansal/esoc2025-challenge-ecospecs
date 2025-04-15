import os
import docx
from typing import List, Any, Optional, Tuple, Dict
from scripts.table_generator import TableGenerator

def test_table_generator():
    generator = TableGenerator(model_path="Qwen/Qwen2.5-7B-Instruct-AWQ")

    row_headers = ["Mercury", "Venus", "Earth", "Mars", "Jupiter"]
    column_headers = ["Distance from Sun (AU)", "Mass (Earth = 1)", "Diameter (km)", "Day length (Earth hours)"]
    
    print("Generating planet data table...")
    table = generator.generate_table_content(
        "Create a table of accurate physical properties for planets in our solar system",
        row_headers=row_headers,
        column_headers=column_headers
    )
    
    # Display the result
    print("\nGenerated Planet Table:")
    for row in table:
        print(row)
    
    # Example 2: Generate a table with auto-generated headers
    print("\nGenerating a table with auto-generated headers...")
    full_table = generator.generate_table(
        "Create a table comparing different programming languages based on popularity and use cases"
    )
    
    # Display the result
    print("\nGenerated Programming Languages Table:")
    for row in full_table:
        print(row)
    
    # Example 3: Save a table to a Word document with intro text
    intro_text = "The following table presents key information about planets in our solar system, including their physical properties and relationship to the Sun."
    
    # Prepare the full table with headers
    planets_full_table = [[""] + column_headers]
    for i, planet in enumerate(row_headers):
        if i < len(table):
            row_data = table[i]
            # Handle various data formats
            if isinstance(row_data, dict):
                # If data is a dictionary, extract values in order of column_headers
                row_values = [row_data.get(col, "N/A") for col in column_headers]
                planets_full_table.append([planet] + row_values)
            elif isinstance(row_data, list):
                # If data is a list, use it directly
                planets_full_table.append([planet] + row_data)
            else:
                # Handle unexpected data types
                planets_full_table.append([planet] + ["N/A"] * len(column_headers))
        else:
            planets_full_table.append([planet] + ["N/A"] * len(column_headers))

    # Save to Word
    generator.save_table_to_docx(
        planets_full_table,
        "generated_planets_table.docx",
        intro_text
    )
    
    # Open and verify the Word document
    doc = docx.Document("generated_planets_table.docx")

    # Check intro text
    assert doc.paragraphs[0].text == intro_text, "Intro text mismatch"

    # Get the table
    table = doc.tables[0]

    # Verify table dimensions
    assert len(table.rows) == len(planets_full_table), "Row count mismatch"
    assert len(table.columns) == len(planets_full_table[0]), "Column count mismatch"

    print("Document verification successful")